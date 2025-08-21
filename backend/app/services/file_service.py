import os
import aiofiles
import PyPDF2
import pdfplumber
from docx import Document
from typing import Dict, Any, Optional
from pathlib import Path
import uuid
import structlog
from ..core.config import settings

logger = structlog.get_logger()

class FileService:
    """文件处理服务"""
    
    def __init__(self):
        self.upload_dir = Path(settings.UPLOAD_DIR)
        self.upload_dir.mkdir(exist_ok=True)
    
    async def save_uploaded_file(self, file_content: bytes, filename: str) -> str:
        """保存上传的文件"""
        try:
            # 生成唯一文件名
            file_extension = Path(filename).suffix
            unique_filename = f"{uuid.uuid4()}{file_extension}"
            file_path = self.upload_dir / unique_filename
            
            # 异步写入文件
            async with aiofiles.open(file_path, 'wb') as f:
                await f.write(file_content)
            
            logger.info(f"文件保存成功", filename=filename, saved_as=unique_filename)
            return str(file_path)
            
        except Exception as e:
            logger.error(f"文件保存失败", filename=filename, error=str(e))
            raise Exception(f"文件保存失败: {str(e)}")
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """从PDF文件提取文本"""
        try:
            text = ""
            
            # 首先尝试使用pdfplumber（更好的文本提取）
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                            
                if text.strip():
                    logger.info(f"使用pdfplumber成功提取PDF文本", length=len(text))
                    return text.strip()
            except Exception as e:
                logger.warning(f"pdfplumber提取失败，尝试PyPDF2", error=str(e))
            
            # 如果pdfplumber失败，尝试PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            
            logger.info(f"使用PyPDF2成功提取PDF文本", length=len(text))
            return text.strip()
            
        except Exception as e:
            logger.error(f"PDF文本提取失败", file_path=file_path, error=str(e))
            raise Exception(f"PDF文本提取失败: {str(e)}")
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """从Word文档提取文本"""
        try:
            doc = Document(file_path)
            text = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
            
            logger.info(f"成功提取Word文档文本", length=len(text))
            return text.strip()
            
        except Exception as e:
            logger.error(f"Word文档文本提取失败", file_path=file_path, error=str(e))
            raise Exception(f"Word文档文本提取失败: {str(e)}")
    
    def extract_text_from_file(self, file_path: str) -> str:
        """根据文件类型提取文本"""
        file_extension = Path(file_path).suffix.lower()
        
        try:
            if file_extension == '.pdf':
                return self.extract_text_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self.extract_text_from_docx(file_path)
            else:
                raise Exception(f"不支持的文件类型: {file_extension}")
                
        except Exception as e:
            logger.error(f"文件文本提取失败", file_path=file_path, error=str(e))
            raise
    
    def validate_file(self, filename: str, file_size: int) -> Dict[str, Any]:
        """验证文件"""
        validation_result = {
            "valid": True,
            "errors": []
        }
        
        # 检查文件大小
        if file_size > settings.MAX_UPLOAD_SIZE:
            validation_result["valid"] = False
            validation_result["errors"].append(
                f"文件大小超过限制 ({file_size} > {settings.MAX_UPLOAD_SIZE})"
            )
        
        # 检查文件扩展名
        file_extension = Path(filename).suffix.lower()
        if file_extension not in settings.ALLOWED_EXTENSIONS:
            validation_result["valid"] = False
            validation_result["errors"].append(
                f"不支持的文件类型: {file_extension}"
            )
        
        return validation_result
    
    async def cleanup_file(self, file_path: str):
        """清理临时文件"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
                logger.info(f"清理临时文件", file_path=file_path)
        except Exception as e:
            logger.warning(f"清理文件失败", file_path=file_path, error=str(e))

# 创建全局文件服务实例
file_service = FileService()